# -*- coding: utf-8 -*-
"""Genera la propuesta comercial en DOCX (estilo profesional)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os

def preparar_logo(src, dst):
    """Recolorea el texto blanco del logo a azul marino para que sea visible sobre fondo blanco."""
    im = Image.open(src).convert('RGBA')
    px = im.load()
    navy = (21, 52, 84)
    for y in range(im.height):
        for x in range(im.width):
            r, g, b, a = px[x, y]
            if a > 0 and r > 225 and g > 225 and b > 225:
                px[x, y] = (navy[0], navy[1], navy[2], a)
    im.save(dst)
    return dst

TEAL   = RGBColor(0x0F, 0x66, 0x74)
ACCENT = RGBColor(0x17, 0xA2, 0xB8)
GRAY   = RGBColor(0x5A, 0x5A, 0x5A)
DARK   = RGBColor(0x22, 0x2A, 0x2E)

doc = Document()

# Márgenes
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.3); s.right_margin = Cm(2.3)

# Fuente base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def bottom_border(p, color="17A2B8", size="12"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom); pPr.append(pbdr)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def kicker(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(2)
    return p

def title(text, sub=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = TEAL
    p.paragraph_format.space_after = Pt(2)
    if sub:
        ps = doc.add_paragraph()
        rs = ps.add_run(sub); rs.font.size = Pt(12); rs.font.color.rgb = GRAY; rs.italic = True
        ps.paragraph_format.space_after = Pt(4)
        bottom_border(ps)

def h1(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(("%s  " % num if num else "") + text)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TEAL
    bottom_border(p, color="BFE6EC", size="6")

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x0D,0x47,0x61)
    return p

def body(text, justify=True):
    p = doc.add_paragraph(text)
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ": "); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

# ─────────────────────────────────────────────── Portada
# Logo MetalCash (recoloreado para fondo blanco), centrado
_logo = preparar_logo('grails-app/assets/images/logo_metalcash_login.png', '/tmp/metalcash_cover.png')
_lp = doc.add_paragraph()
_lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_lp.paragraph_format.space_after = Pt(10)
_lp.add_run().add_picture(_logo, width=Cm(9.0))

kicker("Propuesta Técnica y Comercial")
title("Sistema de Gestión de Liquidación de Minerales",
      "Plataforma web para el ciclo de compra y liquidación de concentrados — Complejo Zinc · Plomo · Plata")

body("La presente propuesta describe una solución de software a la medida para la administración integral "
     "del proceso de compra, control de calidad, valoración y liquidación de concentrados de minerales, "
     "así como el control financiero del proveedor y la generación de reportes operativos y de cumplimiento. "
     "El sistema centraliza la información, automatiza los cálculos y asegura trazabilidad y consistencia en cada etapa.")

# ─────────────────────────────────────────────── 1. Resumen ejecutivo
h1("1.", "Resumen Ejecutivo")
body("El sistema digitaliza el flujo completo desde la recepción del lote hasta el pago al proveedor, "
     "integrando cotizaciones de mercado, precios y términos contractuales, anticipos, retenciones y regalías mineras. "
     "Es una plataforma web multiusuario, con seguridad por roles, cálculos en tiempo real y un servidor que actúa "
     "como única fuente de verdad, eliminando los errores y las inconsistencias propias del manejo en planillas.")
bullets([
    ("Cobertura", "del proceso de recepción, análisis de laboratorio, liquidación, anticipos, transporte y reportes."),
    ("Precisión", "cálculo automático y verificable del valor por tonelada, deducciones, regalías y líquido pagable."),
    ("Control financiero", "estado de cuenta por cliente con manejo de saldos y deudas."),
    ("Cumplimiento", "soporte para regalías mineras, canje de tornaguías y reportes de compras."),
])

# ─────────────────────────────────────────────── 2. Alcance funcional
h1("2.", "Alcance Funcional — Módulos")

h2("2.1  Proceso de Liquidación (núcleo)")
bullets([
    ("Recepción de Lotes", "registro del lote con peso bruto, tara, peso neto húmedo, cantidad de sacos, costo de transporte "
        "y anticipo autorizado; datos de proveedor, empresa, sección, chofer y vehículo; cotizaciones asociadas y numeración por gestión minera."),
    ("Análisis de Laboratorio", "captura de leyes por elemento (zinc, plomo, plata), humedad y merma, con ley registrada y ley de cliente "
        "para soporte de negociación."),
    ("Liquidación", "cálculo automático del peso neto seco, contenidos finos (libras finas y onzas troy), valor bruto de venta, "
        "regalía minera, deducciones y retenciones, bonos, anticipos, líquido pagable y precio. Valoración del mineral por Tabla de Precios "
        "o por Términos de Contrato, recálculo en vivo, comprobante imprimible y anulación controlada."),
])
body("Beneficio: un cálculo exacto, consistente y auditable que reemplaza las planillas manuales, reduce drásticamente los errores "
     "y acelera el cierre de cada liquidación.")

h2("2.2  Anticipos, Pagos y Estado de Cuenta")
bullets([
    ("Anticipo contra Entrega", "anticipos otorgados sobre lotes ya recibidos, con amortización por tramos."),
    ("Anticipo contra Futura Entrega", "anticipos a cuenta de entregas futuras, con comprobante y numeración por gestión."),
    ("Amortización", "registro y seguimiento de la cancelación de anticipos."),
    ("Estado de Cuenta", "libro mayor por cliente (debe / haber / saldo). Un saldo en contra en una liquidación se convierte "
        "automáticamente en deuda, que puede cobrarse —total o en fracciones— en entregas posteriores."),
])
body("Beneficio: control financiero riguroso del proveedor, sin pérdidas de saldos ni dobles cobros.")

h2("2.3  Transporte")
bullets([
    ("Pago por Transporte", "liquidación y registro del servicio de transporte de los lotes."),
    ("Anticipo contra Transporte", "anticipos asociados al servicio de transporte."),
])

h2("2.4  Precios y Condiciones Comerciales")
bullets([
    ("Tablas de Precios", "curvas de ley contra porcentaje pagable por elemento; cálculo del valor por tonelada mediante interpolación lineal, "
        "tomando la cotización diaria del lote y las leyes finales."),
    ("Términos de Contrato", "parámetros contractuales por empresa: deducciones unitarias, porcentajes pagables, maquila, base y escalador, "
        "gastos de refinación y penalidades por impurezas."),
])
body("Beneficio: precios y condiciones totalmente parametrizables, con un valor por tonelada fiel al contrato vigente.")

h2("2.5  Cotizaciones de Mercado")
bullets([
    ("Cotización Diaria y Quincenal", "precios de los minerales utilizados como referencia y base de cálculo."),
    ("Alícuotas", "porcentajes aplicables por elemento."),
    ("Tipo de Cambio (Dólar)", "cotización oficial y comercial para la conversión de valores."),
])

h2("2.6  Retenciones")
bullets([
    ("Configuración de Retenciones", "definición de retenciones por empresa, con distintas bases de cálculo "
        "(sobre valor bruto, valor neto, por saco, por tonelada o monto fijo), aplicadas automáticamente en la liquidación."),
])

h2("2.7  Datos Maestros (Proveedores)")
bullets([
    "Empresas y secciones, Clientes, Choferes, Automóviles y Municipios, como catálogos centralizados que alimentan todo el proceso.",
])

# ─────────────────────────────────────────────── 3. Reportes
h1("3.", "Reportes Operativos y de Cumplimiento")
body("El sistema contempla un conjunto de reportes para la gestión operativa, contable y de cumplimiento normativo, "
     "con posibilidad de impresión y exportación:")
bullets([
    ("Lotes Recepcionados", "detalle de los lotes ingresados en un período."),
    ("Lotes y Análisis", "lotes con sus resultados de laboratorio."),
    ("Lotes Liquidados", "liquidaciones realizadas con sus valores."),
    ("Planilla de Liquidación", "consolidado de liquidaciones para control y pago."),
    ("Detalle de Compras", "resumen de las compras de mineral por período."),
    ("Canje de Tornaguías", "control del canje de tornaguías de transporte."),
    ("Libro de Regalías Mineras (RM)", "registro de las regalías mineras de las compras, para cumplimiento."),
    ("Anticipos contra Entrega", "seguimiento de los anticipos otorgados."),
    ("Pago de Transporte", "reporte de los pagos por transporte."),
    ("Retenciones", "detalle de las retenciones aplicadas."),
    ("Gráfico de Cantidad", "visualización gráfica de volúmenes para análisis gerencial."),
])
body("Beneficio: información oportuna para la toma de decisiones y respaldo documental para auditoría y obligaciones ante las autoridades mineras.")

# ─────────────────────────────────────────────── 4. Características técnicas
h1("4.", "Características Transversales")
bullets([
    ("Plataforma web", "acceso desde el navegador, con interfaz moderna y adaptable a distintos dispositivos."),
    ("Seguridad por roles", "perfiles diferenciados (Recepción, Control de Calidad, Liquidación, Caja, Reportes y Administración)."),
    ("Cálculo en tiempo real", "los valores se recalculan al instante; el servidor valida y recalcula como fuente de verdad."),
    ("Trazabilidad", "numeración por gestión minera, registro de movimientos y anulaciones controladas."),
    ("Impresión profesional", "comprobantes de liquidación y reportes ejecutivos listos para imprimir (formato Carta)."),
    ("Base de datos relacional", "información estructurada, íntegra y respaldable."),
])

# ─────────────────────────────────────────────── 5. Beneficios
h1("5.", "Beneficios para su Empresa")
bullets([
    "Reducción drástica de errores de cálculo respecto al manejo en hojas de cálculo.",
    "Mayor velocidad en el cierre de liquidaciones y pagos.",
    "Control financiero y de deudas por proveedor.",
    "Cumplimiento y respaldo documental (regalías, tornaguías, compras).",
    "Información centralizada y reportes para la toma de decisiones.",
    "Solución escalable, adaptable al crecimiento y a nuevos requerimientos.",
])

# ─────────────────────────────────────────────── 6. Inversión
h1("6.", "Inversión Estimada")
body("La inversión estimada para la implementación de la solución descrita es la siguiente:")

tbl = doc.add_table(rows=2, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "Concepto"; hdr[1].text = "Monto (Bs)"
for c in hdr:
    shade(c, "0F6674")
    for pr in c.paragraphs:
        for rn in pr.runs:
            rn.bold = True; rn.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
row = tbl.rows[1].cells
row[0].text = "Implementación del Sistema de Gestión de Liquidación de Minerales (estimado)"
row[1].text = "5.000,00"
row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
nota = doc.add_paragraph()
rn = nota.add_run("Monto referencial expresado en bolivianos (Bs). El alcance definitivo y las condiciones "
                  "de pago se acordarán de común acuerdo según los requerimientos finales.")
rn.italic = True; rn.font.size = Pt(9); rn.font.color.rgb = GRAY

# ─────────────────────────────────────────────── Cierre
h1("7.", "Cierre")
body("Agradecemos la oportunidad de presentar esta propuesta. La solución está diseñada para ordenar, agilizar y dar "
     "transparencia a su operación de compra y liquidación de minerales. Quedamos a su disposición para una demostración "
     "y para ajustar el alcance a las necesidades específicas de su empresa.")

# Pie de página
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Sistema de Gestión de Liquidación de Minerales  ·  Propuesta confidencial")
fr.font.size = Pt(8); fr.font.color.rgb = GRAY

# Propiedades
cp = doc.core_properties
cp.title = "Propuesta — Sistema de Gestión de Liquidación de Minerales"
cp.subject = "Propuesta técnica y comercial"
cp.category = "Propuesta"

out = "Propuesta_Sistema_Liquidaciones.docx"
doc.save(out)
print("OK:", out)
